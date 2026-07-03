from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from agentic_profile_matching import config

RESUMES = [
    {
        "name": "Alex Kumar",
        "role": "Machine Learning Engineer",
        "filename": "resume_alex_kumar.txt",
        "exp": 4,
        "education": "B.Tech in Computer Science, IIT",
        "skills": ["Python", "Machine Learning", "LangChain", "RAG", "Data Science", "PyTorch", "Git"],
        "summary": "Passionate Machine Learning Engineer specializing in NLP, Large Language Models, and RAG systems.",
        "experience": "Machine Learning Engineer at TechAI (2022-Present)\n- Built advanced RAG systems using LangChain and Python.\n- Improved retrieval accuracy by 25% using hybrid search.\n\nData Scientist at DataCorp (2020-2022)\n- Analyzed user data using Pandas and NumPy.\n- Built predictive models using Scikit-Learn."
    },
    {
        "name": "Emily Watson",
        "role": "Senior Backend Engineer",
        "filename": "resume_emily_watson.docx",
        "exp": 8,
        "education": "M.S. in Software Engineering, Stanford",
        "skills": ["Python", "Django", "FastAPI", "PostgreSQL", "Docker", "Kubernetes", "AWS", "CI/CD"],
        "summary": "Senior Backend Engineer with 8+ years of experience building scalable microservices and database-driven web applications.",
        "experience": "Senior Backend Engineer at CloudTech (2021-Present)\n- Designed microservices architecture handling 10k+ requests per second.\n- Orchestrated containers using Docker and Kubernetes.\n\nSoftware Engineer at WebSystems (2018-2021)\n- Developed backend APIs using Python and Django."
    },
    {
        "name": "Michael Lee",
        "role": "DevOps & Cloud Engineer",
        "filename": "resume_michael_lee.pdf",
        "exp": 6,
        "education": "B.S. in Computer Engineering, Berkeley",
        "skills": ["AWS", "Docker", "Kubernetes", "Terraform", "CI/CD", "Jenkins", "Linux", "Python"],
        "summary": "DevOps Engineer focused on automating infrastructure, optimizing CI/CD pipelines, and cloud migration.",
        "experience": "Cloud Infrastructure Engineer at ScaleOps (2020-Present)\n- Managed AWS infrastructure using Terraform (Infrastructure as Code).\n- Implemented robust CI/CD pipelines utilizing Jenkins and GitHub Actions.\n\nSystems Administrator at TechBase (2020-2020)\n- Maintained Linux servers and monitored uptime."
    },
    {
        "name": "Shashank Chandel",
        "role": "Distributed Systems Engineer",
        "filename": "resume_shashank_chandel.txt",
        "exp": 5,
        "education": "B.Tech in Information Technology, SRMIST",
        "skills": ["Java", "Spring Boot", "Kafka", "SQL", "Python", "Generative AI", "Docker", "Kubernetes", "AWS"],
        "summary": "Distributed systems engineer with 5+ years of backend engineering experience building scalable microservices.",
        "experience": "Backend Engineer at KafkaTech (2021-Present)\n- Engineered highly concurrent streaming pipelines using Kafka and Spring Boot.\n- Deployed production Docker containers to AWS EKS.\n\nSoftware Developer at DataSoft (2021-2021)\n- Developed Java backend services and REST APIs."
    },
    {
        "name": "Jane Doe",
        "role": "Frontend Developer",
        "filename": "resume_jane_doe.docx",
        "exp": 3,
        "education": "B.S. in Computer Science, NYU",
        "skills": ["JavaScript", "TypeScript", "React", "HTML", "CSS", "GraphQL", "Git"],
        "summary": "Frontend developer specializing in building modern, responsive user interfaces and single-page apps using React and TypeScript.",
        "experience": "Frontend Developer at DesignCorp (2023-Present)\n- Re-architected company website in React and TypeScript.\n- Integrated GraphQL endpoints with frontend components.\n\nJunior Web Developer at Pixels (2023-2023)\n- Developed responsive HTML and CSS pages."
    },
    {
        "name": "John Doe",
        "role": "Backend Developer",
        "filename": "resume_john_doe.pdf",
        "exp": 5,
        "education": "M.S. in Computer Science, Georgia Tech",
        "skills": ["Python", "FastAPI", "SQL", "PostgreSQL", "Docker", "Git", "REST API"],
        "summary": "Backend Developer with 5 years of experience building APIs, managing databases, and containerizing software.",
        "experience": "Backend Developer at AppHouse (2021-Present)\n- Created REST APIs using FastAPI and Python.\n- Managed PostgreSQL database performance.\n\nSoftware Developer at StartUp (2021-2021)\n- Maintained and optimized database queries."
    },
    {
        "name": "Sarah Connor",
        "role": "Data Scientist",
        "filename": "resume_sarah_connor.txt",
        "exp": 7,
        "education": "PhD in Statistics, MIT",
        "skills": ["Python", "Data Science", "Machine Learning", "Deep Learning", "TensorFlow", "Pandas", "NumPy", "SQL"],
        "summary": "Data Scientist with a solid academic background in Statistics and expertise in deep learning, predictive modeling, and big data analysis.",
        "experience": "Lead Data Scientist at Cyberdyne (2019-Present)\n- Developed deep learning models using TensorFlow for predictive maintenance.\n- Analyzed complex datasets using Pandas and NumPy."
    },
    {
        "name": "Kyle Reese",
        "role": "Mobile iOS Developer",
        "filename": "resume_kyle_reese.docx",
        "exp": 4,
        "education": "B.S. in Computer Science, UCLA",
        "skills": ["Swift", "Objective-C", "iOS", "Flutter", "React Native", "Git", "REST API"],
        "summary": "Mobile App Developer focused on creating elegant iOS applications and cross-platform mobile experiences.",
        "experience": "iOS Developer at AppStudio (2022-Present)\n- Developed native iOS apps in Swift.\n- Built cross-platform mobile apps with Flutter."
    },
    {
        "name": "David Bowman",
        "role": "Cloud Solutions Architect",
        "filename": "resume_david_bowman.pdf",
        "exp": 9,
        "education": "M.S. in Distributed Systems, UIUC",
        "skills": ["GCP", "Azure", "AWS", "Kubernetes", "Docker", "Linux", "Terraform", "CI/CD"],
        "summary": "Cloud Solutions Architect with 9+ years of experience in multi-cloud strategies, container orchestration, and serverless architectures.",
        "experience": "Solutions Architect at Hal9000 Cloud (2017-Present)\n- Designed enterprise architectures across GCP and AWS.\n- Orchestrated large-scale Kubernetes deployments."
    },
    {
        "name": "Ellen Ripley",
        "role": "Engineering Manager",
        "filename": "resume_ellen_ripley.txt",
        "exp": 12,
        "education": "M.S. in Engineering Management, Cornell",
        "skills": ["Git", "CI/CD", "Linux", "AWS", "Python", "SQL", "Docker"],
        "summary": "Engineering Manager with 12+ years of technical leadership experience managing cross-functional software teams.",
        "experience": "Engineering Manager at Weyland-Yutani (2014-Present)\n- Managed teams of 15+ backend and DevOps engineers.\n- Directed migration of legacy systems to AWS."
    },
    {
        "name": "Marcus Wright",
        "role": "Full Stack Developer",
        "filename": "resume_marcus_wright.docx",
        "exp": 3,
        "education": "B.S. in Computer Science, UT Austin",
        "skills": ["JavaScript", "TypeScript", "React", "Node.js", "Express", "MongoDB", "CSS", "HTML"],
        "summary": "Full Stack Engineer specializing in the MERN stack with 3 years of experience building web applications.",
        "experience": "Full Stack Developer at ProjectAngel (2023-Present)\n- Created user-facing features using React.\n- Designed backend services using Node.js and Express."
    },
    {
        "name": "John Connor",
        "role": "Rust Systems Developer",
        "filename": "resume_john_connor.pdf",
        "exp": 4,
        "education": "B.Tech in Computer Science, IIT Delhi",
        "skills": ["Rust", "C++", "Linux", "Git", "Docker", "SQL", "WebAssembly"],
        "summary": "Systems Programmer specializing in memory-safe performance code using Rust and WebAssembly.",
        "experience": "Systems Developer at Resistance Tech (2022-Present)\n- Engineered low-latency backend systems using Rust.\n- Managed compile-time safety and profiling tools."
    },
    {
        "name": "Peter Parker",
        "role": "Python Developer",
        "filename": "resume_peter_parker.txt",
        "exp": 2,
        "education": "B.S. in Computer Science, Queens College",
        "skills": ["Python", "Flask", "Django", "SQL", "Git", "HTML", "CSS"],
        "summary": "Junior Python Developer with experience in web scraping, scripting, and backend web development using Flask and Django.",
        "experience": "Python Developer at DailyBugle (2024-Present)\n- Built web scrapers and internal data tools using Python.\n- Maintained database records using SQL."
    },
    {
        "name": "Bruce Wayne",
        "role": "Security Engineer",
        "filename": "resume_bruce_wayne.docx",
        "exp": 10,
        "education": "M.S. in Cybersecurity, Johns Hopkins",
        "skills": ["Linux", "Git", "Python", "Docker", "Kubernetes", "AWS", "CI/CD"],
        "summary": "Cybersecurity Specialist focused on penetration testing, security automation, and secure infrastructure design.",
        "experience": "Security Architect at Wayne Enterprises (2016-Present)\n- Designed secure network infrastructures and led vulnerability scans.\n- Automated security policies in AWS using Terraform and Ansible."
    },
    {
        "name": "Clark Kent",
        "role": "Data Engineer",
        "filename": "resume_clark_kent.pdf",
        "exp": 5,
        "education": "B.S. in Computer Science, Metropolis University",
        "skills": ["Python", "Spark", "Hadoop", "SQL", "PostgreSQL", "Kafka", "AWS", "Git"],
        "summary": "Data Engineer specializing in constructing scalable ETL data pipelines and managing large-scale data warehouses.",
        "experience": "Data Engineer at Planet Media (2021-Present)\n- Built real-time streaming pipelines using Kafka and Spark.\n- Created scalable ETL jobs running on AWS."
    },
    {
        "name": "Diana Prince",
        "role": "Product Manager",
        "filename": "resume_diana_prince.txt",
        "exp": 8,
        "education": "MBA, Harvard Business School",
        "skills": ["Git", "Python", "SQL", "Jira", "AWS"],
        "summary": "Technical Product Manager with a strong software engineering background, managing roadmap and release cycles for cloud products.",
        "experience": "Senior PM at Olympus Tech (2018-Present)\n- Launched three major SaaS products generating $5M+ ARR.\n- Wrote technical specifications and query databases using SQL."
    },
    {
        "name": "Barry Allen",
        "role": "Golang Engineer",
        "filename": "resume_barry_allen.docx",
        "exp": 3,
        "education": "B.S. in Computer Science, Central City University",
        "skills": ["Go", "Golang", "SQL", "PostgreSQL", "Docker", "REST API", "Git", "Redis"],
        "summary": "Golang developer focused on constructing high-throughput backend services and concurrent microservices.",
        "experience": "Go Developer at StarLabs (2023-Present)\n- Developed scalable API handlers and middleware in Go.\n- Optimized Redis caching layers for quick queries."
    },
    {
        "name": "Hal Jordan",
        "role": "DevOps Engineer",
        "filename": "resume_hal_jordan.pdf",
        "exp": 6,
        "education": "B.S. in Systems Engineering, Ferris Air Academy",
        "skills": ["AWS", "Terraform", "CI/CD", "Ansible", "Linux", "Docker", "Git"],
        "summary": "DevOps Engineer specializing in continuous integration and automated deployment configurations.",
        "experience": "DevOps Engineer at Sector2814 (2020-Present)\n- Engineered robust build systems and testing automation pipelines.\n- Configured automated servers using Ansible and Terraform."
    },
    {
        "name": "Arthur Curry",
        "role": "Database Administrator",
        "filename": "resume_arthur_curry.txt",
        "exp": 7,
        "education": "M.S. in Database Systems, University of Maine",
        "skills": ["SQL", "PostgreSQL", "MySQL", "MongoDB", "Cassandra", "Redis", "Linux"],
        "summary": "Database Administrator specializing in high-availability database architectures, replication, and query tuning.",
        "experience": "Lead DBA at AtlantisData (2019-Present)\n- Maintained large PostgreSQL and Cassandra clusters.\n- Spearheaded database replication and backup strategy."
    },
    {
        "name": "Victor Stone",
        "role": "AI Research Scientist",
        "filename": "resume_victor_stone.docx",
        "exp": 5,
        "education": "PhD in Computer Science, Caltech",
        "skills": ["Python", "PyTorch", "TensorFlow", "Deep Learning", "Machine Learning", "NLP", "C++"],
        "summary": "AI Scientist researching deep learning architectures, computer vision, and neural network optimization.",
        "experience": "Research Scientist at S.T.A.R. Labs (2021-Present)\n- Published papers on neural network optimization.\n- Developed custom PyTorch neural layers."
    },
    {
        "name": "Tony Stark",
        "role": "Solutions Architect",
        "filename": "resume_tony_stark.pdf",
        "exp": 15,
        "education": "PhD in Physics and Engineering, MIT",
        "skills": ["C++", "Python", "AWS", "Docker", "Kubernetes", "Linux", "Microservices", "Git"],
        "summary": "IoT Solutions Architect with a history of building highly-secure distributed systems and smart automation.",
        "experience": "Chief Architect at Stark Industries (2011-Present)\n- Designed smart grid IoT architectures and microservices.\n- Managed high-performance compute clusters."
    },
    {
        "name": "Steve Rogers",
        "role": "Scrum Master",
        "filename": "resume_steve_rogers.txt",
        "exp": 8,
        "education": "B.S. in Business Administration, Brooklyn College",
        "skills": ["Git", "Jira", "Linux", "SQL", "Python"],
        "summary": "Certified Scrum Master with 8 years of experience leading agile transitions and managing software development lifecycles.",
        "experience": "Scrum Master at ShieldTech (2018-Present)\n- Facilitated agile ceremonies and sprint planning for 4 developers.\n- Tracked velocity and release metrics."
    },
    {
        "name": "Natasha Romanoff",
        "role": "QA Automation Engineer",
        "filename": "resume_natasha_romanoff.docx",
        "exp": 6,
        "education": "B.S. in Software Testing, Kiev University",
        "skills": ["Python", "Selenium", "Git", "CI/CD", "Linux", "SQL", "PostgreSQL"],
        "summary": "QA Engineer specializing in building automated test suites, functional testing, and integration testing for web platforms.",
        "experience": "QA Automation Engineer at RedTech (2020-Present)\n- Created automated Selenium testing frameworks using Python.\n- Integrated QA suites into CI/CD pipelines."
    },
    {
        "name": "Bruce Banner",
        "role": "Bioinformatics Scientist",
        "filename": "resume_bruce_banner.pdf",
        "exp": 10,
        "education": "PhD in Biochemistry & CS, Berkeley",
        "skills": ["Python", "Pandas", "NumPy", "Data Science", "SQL", "Linux", "Git"],
        "summary": "Bioinformatics scientist analyzing genomic sequences and building big data pipelines for biomedical startups.",
        "experience": "Research Fellow at GammaGen (2016-Present)\n- Developed python-based sequence analysis algorithms.\n- Optimized large dataset analysis with Pandas and NumPy."
    },
    {
        "name": "Clint Barton",
        "role": "Frontend React Developer",
        "filename": "resume_clint_barton.txt",
        "exp": 4,
        "education": "B.S. in Web Design, Oregon State",
        "skills": ["React", "JavaScript", "TypeScript", "HTML5", "CSS3", "Git", "GraphQL"],
        "summary": "Frontend developer focused on building interactive, user-friendly, and lightweight web apps in React.",
        "experience": "Frontend Developer at ArrowApps (2022-Present)\n- Developed responsive frontend layouts in React.\n- Styled interfaces using CSS3 and Tailwind CSS."
    },
    {
        "name": "Wanda Maximoff",
        "role": "ML Ops Engineer",
        "filename": "resume_wanda_maximoff.docx",
        "exp": 5,
        "education": "M.S. in Computational Linguistics, Sokovia University",
        "skills": ["Python", "Docker", "Kubernetes", "AWS", "Terraform", "CI/CD", "Machine Learning", "Git"],
        "summary": "MLOps Engineer establishing automated machine learning model deployment pipelines and monitoring infrastructures.",
        "experience": "MLOps Engineer at MagicAI (2021-Present)\n- Containerized ML models using Docker and Kubernetes.\n- Created automation scripts in Python."
    },
    {
        "name": "Vision",
        "role": "Compiler Systems Engineer",
        "filename": "resume_vision.pdf",
        "exp": 8,
        "education": "PhD in Computer Architecture, Oxford",
        "skills": ["C++", "Rust", "LLVM", "Linux", "Git", "Docker", "WebAssembly"],
        "summary": "Compiler engineer with extensive knowledge of programming language internals and binary optimization techniques.",
        "experience": "Compiler Engineer at MindSoft (2018-Present)\n- Wrote code optimization passes inside LLVM.\n- Maintained compiler frontends using C++ and Rust."
    },
    {
        "name": "Sam Wilson",
        "role": "Cloud Security Specialist",
        "filename": "resume_sam_wilson.txt",
        "exp": 5,
        "education": "B.S. in Cybersecurity, Howard University",
        "skills": ["AWS", "Azure", "Linux", "Git", "CI/CD", "Docker", "Terraform"],
        "summary": "Cloud Security Consultant analyzing cloud architecture vulnerabilities and establishing strict IAM security configurations.",
        "experience": "Cloud Security Engineer at FalconCloud (2021-Present)\n- Configured cloud firewalls and IAM roles on AWS.\n- Conducted container scanning and Docker auditing."
    },
    {
        "name": "Bucky Barnes",
        "role": "C++ Systems Programmer",
        "filename": "resume_bucky_barnes.docx",
        "exp": 9,
        "education": "B.S. in Computer Science, Rutgers",
        "skills": ["C++", "Linux", "Git", "Docker", "SQL", "Rest API"],
        "summary": "Low-level C++ engineer with experience writing kernel modules, driver software, and multithreaded systems.",
        "experience": "Systems Developer at WinterTech (2017-Present)\n- Programmed backend systems and multithreaded apps in C++.\n- Debugged complex software loops using GDB."
    },
    {
        "name": "Carol Danvers",
        "role": "Aerospace Software Engineer",
        "filename": "resume_carol_danvers.pdf",
        "exp": 11,
        "education": "M.S. in Aeronautics and Computer Science, MIT",
        "skills": ["C++", "Python", "Linux", "Git", "Docker", "RTOS", "Embedded Systems"],
        "summary": "Embedded software engineer writing safety-critical software and real-time operating system applications.",
        "experience": "Embedded Lead at Starforce (2015-Present)\n- Developed safety-critical code in C++ for aerospace equipment.\n- Integrated RTOS software with custom boards."
    },
    {
        "name": "Peter Quill",
        "role": "Mobile React Native Developer",
        "filename": "resume_peter_quill.txt",
        "exp": 5,
        "education": "B.S. in Media Arts, Colorado",
        "skills": ["React Native", "React", "JavaScript", "TypeScript", "Git", "HTML", "CSS"],
        "summary": "Mobile software engineer crafting music apps and rich interactive media apps using React Native.",
        "experience": "Mobile App Engineer at MilanoApps (2021-Present)\n- Published four cross-platform apps on Google Play Store.\n- Designed customized responsive screens in React Native."
    }
]

JOB_DESCRIPTIONS = [
    {
        "id": "jd_python_ml.txt",
        "content": "Job Description: Python & Machine Learning Engineer\n\nWe are looking for a Machine Learning Engineer to join our team. The candidate will build production NLP systems, RAG models, and large language model systems.\nRequirements:\n- Must have 3+ years of experience\n- Strong experience in Python, Machine Learning, and libraries like PyTorch/TensorFlow\n- Experience with LangChain, RAG systems, and generative AI is a strong plus\n- Knowledge of Git and software engineering practices."
    },
    {
        "id": "jd_java_spring.txt",
        "content": "Job Description: Senior Java & Spring Boot Developer\n\nWe are hiring a Senior Java Developer to build distributed backend systems.\nRequirements:\n- Must have 6+ years of experience\n- Advanced skills in Java and Spring Boot framework\n- Experience building streaming pipelines with Kafka\n- Strong SQL knowledge\n- Experience with container tools (Docker, Kubernetes) is a plus."
    },
    {
        "id": "jd_devops_cloud.txt",
        "content": "Job Description: DevOps & Cloud Infrastructure Architect\n\nWe are looking for a DevOps Architect to automate deployments and manage cloud environments.\nRequirements:\n- Must have 5+ years of experience\n- Expert knowledge of AWS cloud provider\n- Strong experience with Terraform (Infrastructure as Code)\n- Experience with container orchestration using Docker and Kubernetes\n- Strong CI/CD automation skill using Jenkins or GitHub Actions."
    },
    {
        "id": "jd_frontend_react.txt",
        "content": "Job Description: Frontend Web Developer (React & TypeScript)\n\nWe need a Frontend Developer to build clean, interactive web user interfaces.\nRequirements:\n- Must have 3+ years of experience\n- Expert skills in React, JavaScript, and TypeScript\n- Strong knowledge of HTML5, CSS3, and responsive design\n- Experience with GraphQL APIs and state management."
    },
    {
        "id": "jd_fullstack_go.txt",
        "content": "Job Description: Full Stack Developer (Golang & React)\n\nWe are looking for a developer to work on both backend services and frontend applications.\nRequirements:\n- Must have 4+ years of experience\n- Solid experience building backend APIs in Go / Golang\n- Experience with React frontend development\n- Good knowledge of SQL (PostgreSQL / MySQL)\n- Familiarity with Docker and Git."
    }
]


def generate_text_resume(candidate, path):
    content = f"""{candidate['name']}
{candidate['role']}

SUMMARY:
{candidate['summary']}

SKILLS:
{', '.join(candidate['skills'])}

EXPERIENCE:
{candidate['experience']}

EDUCATION:
{candidate['education']}
"""
    Path(path).write_text(content, encoding='utf-8')


def generate_docx_resume(candidate, path):
    doc = Document()
    doc.add_heading(candidate['name'], level=0)
    doc.add_heading(candidate['role'], level=1)
    
    doc.add_heading('SUMMARY', level=2)
    doc.add_paragraph(candidate['summary'])
    
    doc.add_heading('SKILLS', level=2)
    doc.add_paragraph(', '.join(candidate['skills']))
    
    doc.add_heading('EXPERIENCE', level=2)
    for line in candidate['experience'].split('\n'):
        doc.add_paragraph(line)
        
    doc.add_heading('EDUCATION', level=2)
    doc.add_paragraph(candidate['education'])
    
    doc.save(path)


def generate_pdf_resume(candidate, path):
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=letter)
    story = []
    
    story.append(Paragraph(candidate['name'], styles['Title']))
    story.append(Paragraph(candidate['role'], styles['Heading2']))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph('SUMMARY', styles['Heading3']))
    story.append(Paragraph(candidate['summary'], styles['BodyText']))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph('SKILLS', styles['Heading3']))
    story.append(Paragraph(', '.join(candidate['skills']), styles['BodyText']))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph('EXPERIENCE', styles['Heading3']))
    for line in candidate['experience'].split('\n'):
        if line.strip():
            story.append(Paragraph(line, styles['BodyText']))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph('EDUCATION', styles['Heading3']))
    story.append(Paragraph(candidate['education'], styles['BodyText']))
    
    doc.build(story)


def main():
    resumes_dir = Path(config.RESUMES_DIR)
    jds_dir = Path(config.JOB_DESCRIPTIONS_DIR)
    
    resumes_dir.mkdir(parents=True, exist_ok=True)
    jds_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"Generating {len(RESUMES)} resumes in {resumes_dir}...")
    for c in RESUMES:
        # Prepend experience to summary
        c['summary'] = f"{c['role']} with {c['exp']}+ years of experience. {c['summary']}"
        
        filepath = resumes_dir / c['filename']
        ext = filepath.suffix.lower()
        if ext == '.txt':
            generate_text_resume(c, str(filepath))
        elif ext == '.docx':
            generate_docx_resume(c, str(filepath))
        elif ext == '.pdf':
            generate_pdf_resume(c, str(filepath))
        print(f"Generated: {c['filename']}")
        
    print(f"Generating {len(JOB_DESCRIPTIONS)} job descriptions in {jds_dir}...")
    for jd in JOB_DESCRIPTIONS:
        filepath = jds_dir / jd['id']
        filepath.write_text(jd['content'], encoding='utf-8')
        print(f"Generated: {jd['id']}")

    print("Dataset generation complete.")


if __name__ == "__main__":
    main()
